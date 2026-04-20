import sys, subprocess, re

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "--quiet", "python-docx"])
    import docx

from docx import Document

doc = Document()

md_file = r"C:\Users\Asus\.gemini\antigravity\brain\d519475e-526d-42be-be99-2392bc1b5525\artifacts\project_report.md"
with open(md_file, "r", encoding="utf-8") as f:
    lines = f.readlines()

for line in lines:
    line = line.strip()
    if not line or line.startswith('---'):
        continue
    
    if line.startswith('<h1') or line.startswith('</h1>'):
        continue
        
    if line.startswith('# '):
        doc.add_heading(line[2:].replace('**', ''), level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:].replace('**', ''), level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:].replace('**', ''), level=3)
    elif line.startswith('#### '):
        doc.add_heading(line[5:].replace('**', ''), level=4)
    elif line.startswith('- ') or line.startswith('* '):
        p = doc.add_paragraph('', style='List Bullet')
        text = line[2:]
        parts = text.split('**')
        for i, part in enumerate(parts):
            p.add_run(part).bold = (i % 2 != 0)
    elif re.match(r'^\d+\.', line):
        p = doc.add_paragraph('', style='List Number')
        text = line.split('.', 1)[1].strip()
        parts = text.split('**')
        for i, part in enumerate(parts):
            p.add_run(part).bold = (i % 2 != 0)
    elif line.startswith('> '):
        doc.add_paragraph(line[2:], style='Quote')
    elif line.startswith('![') or line.startswith('*(Note:'):
        p = doc.add_paragraph(line)
        p.alignment = 1 # centered
    else:
        p = doc.add_paragraph('')
        parts = line.split('**')
        for i, part in enumerate(parts):
            p.add_run(part).bold = (i % 2 != 0)

save_path = r"C:\Web App Development\Kaizen ERM\Kaizen_ERM_Project_Report.docx"
doc.save(save_path)
print(f"SUCCESS: DOCX created at: {save_path}")
